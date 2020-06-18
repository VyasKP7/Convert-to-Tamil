#!/usr/bin/env python3

import tkinter as tk
from tkinter import filedialog, Text
import docx

root = tk.Tk()
root.title("Convert Between Tamil and English")


def get_text(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def put_text(filename, inp):
    doc = docx.Document()
    doc.add_paragraph(inp)
    name = filename[:-4] + '.docx'
    doc.save(name)
    my_label = tk.Label(root, text="File named {a} added to folder".format(a=name))
    my_label.pack()


def put_txt(filename, inp):
    #Creates a txt File
    nhandle = open(filename, 'w')
    nhandle.write(inp)
    my_label = tk.Label(root, text="File named {a} added to folder".format(a=filename))
    my_label.pack()


def put_txtb(filename,inp):
    #creates a Docx file
    nhandle = open(filename, 'wb')
    nhandle.write(inp.encode('utf8'))
    my_label = tk.Label(root, text= "File named {a} added to folder".format(a=filename))
    my_label.pack()

def tamil_replace(lyrics, dir):
    engtotam = {
        "saa": "ஸா",
        "sa": "ஸ",
        "ree": "ரீ",
        "ri": "ரி",
        "gaa": "கா",
        "ga": "க",
        "maa": "மா",
        "ma": "ம",
        "paa": "பா",
        "pa": "ப",
        "daa": "தா",
        "da": "த",
        "nii": "நீ",
        "ni": "நி"
    }
    lyrics = lyrics.lower()
    if dir == 'tamil':
        for eng, tam in engtotam.items():
            lyrics = lyrics.replace(eng, tam)
    else:
        for eng, tam in engtotam.items():
            lyrics = lyrics.replace(tam, eng)
    return lyrics

def tamil_convert():
    filename = filedialog.askopenfilename(initialdir="/", title="Select File",
                                          filetypes = (("text", "*.txt"),("Word", "*.docx")))
    flag = 0
    if filename.endswith(".txt"):
        fhandle = open(filename, 'r')
        lyrics = fhandle.read()
        lyrics = tamil_replace(lyrics, 'tamil')
    else:
        flag = 1
        lyrics = get_text(filename)
        lyrics = tamil_replace(lyrics, 'tamil')

    i = filename.rindex('/')
    if flag == 1:
        name = filename[i + 1:-5]
    else:
        name = filename[i + 1:-4]

    directory = filename[:i + 1]
    name = name + '(tamil).txt'
    nname = directory + name

    if flag == 1:
        put_text(nname, lyrics)
    else:
        put_txtb(nname, lyrics)


def english_convert():
    filename = filedialog.askopenfilename(initialdir="/", title="Select File",
                                          filetypes = (("Word", "*.docx"), ("text", "*.txt")))
    flag = 0
    if filename.endswith(".txt"):
        fhandle = open(filename, 'rb')
        lyrics = fhandle.read()
        lyrics = lyrics.decode()
        lyrics = tamil_replace(lyrics, 'english')
    else:
        flag = 1
        lyrics = get_text(filename)
        lyrics = tamil_replace(lyrics, 'english')


    i = filename.rindex('/')

    if flag == 1:
        name = filename[i + 1:-5]
    else:
        name = filename[i + 1:-4]

    directory = filename[:i + 1]
    name = name + '(tamil).txt'
    nname = directory + name

    if flag == 1:
        put_text(nname, lyrics)
    else:
        put_txt(nname,lyrics)



myButton1 = tk.Button(root, text="Convert file from English to Tamil", command=tamil_convert)
myButton1.pack()
myButton1 = tk.Button(root, text="Convert file from Tamil to English", command=english_convert)
myButton1.pack()
root.mainloop()

